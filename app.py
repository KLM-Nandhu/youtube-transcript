import streamlit as st
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pytube import YouTube
import re
import os
import base64
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import io

# Custom CSS for buttons, text, background, and animation
st.markdown("""
<style>
    body {
        background: linear-gradient(135deg, #e6f2ff 0%, #ffffff 100%);
        background-attachment: fixed;
    }
    .stApp {
        background: transparent;
    }
    .stButton>button {
        width: auto;
        background-color: #4a90e2;
        color: white;
        border: none;
        padding: 0.75rem 1.5rem;
        font-size: 1.1rem;
        cursor: pointer;
        transition: background-color 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #357abd;
    }
    .big-font {
        font-size: 20px !important;
        color: #2c3e50;
    }
    .medium-font {
        font-size: 16px !important;
        color: #34495e;
    }
    .video-info {
        display: flex;
        align-items: center;
        margin-bottom: 20px;
        background-color: rgba(255, 255, 255, 0.7);
        padding: 10px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    .video-info img {
        width: 120px;
        height: 90px;
        margin-right: 20px;
        border-radius: 5px;
    }
    .video-title {
        font-size: 18px;
        font-weight: bold;
        color: #2c3e50;
    }
    .title-container {
        display: flex;
        justify-content: center;
        margin-left: 100px;
        margin-bottom: 30px;
    }
    .animated-text {
        position: fixed;
        top: 70px;
        left: 20px;
        font-size: 28px;
        font-weight: bold;
        color: #2c3e50;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
        animation: softMoveAndGlow 3s ease-in-out infinite;
    }
    @keyframes softMoveAndGlow {
        0% { transform: translateY(0px); text-shadow: 2px 2px 4px rgba(0,0,0,0.1); }
        50% { transform: translateY(-10px); text-shadow: 0 0 10px rgba(74, 144, 226, 0.5); }
        100% { transform: translateY(0px); text-shadow: 2px 2px 4px rgba(0,0,0,0.1); }
    }
    .stTextInput>div>div>input {
        font-size: 16px;
    }
    .background-design {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background-image: 
            radial-gradient(circle at 10% 20%, rgba(74, 144, 226, 0.1) 0%, transparent 50%),
            radial-gradient(circle at 90% 80%, rgba(74, 144, 226, 0.1) 0%, transparent 50%);
        z-index: -1;
    }
</style>

<div class="background-design"></div>
<div class="animated-text">Bent's Woodworking</div>
""", unsafe_allow_html=True)

def get_video_info(video_id):
    try:
        url = f"https://www.youtube.com/watch?v={video_id}"
        yt = YouTube(url)
        return yt.title, yt.thumbnail_url
    except Exception as e:
        return f"An error occurred while fetching video info: {str(e)}", None

def get_video_transcript_with_timestamps(video_id):
    try:
        transcript = YouTubeTranscriptApi.get_transcript(video_id)
        return transcript
    except Exception as e:
        return f"An error occurred while fetching the transcript: {str(e)}"

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "", filename).replace(" ", "_")

def save_transcript_to_word_and_pdf(video_id):
    title, _ = get_video_info(video_id)
    if isinstance(title, str) and title.startswith("An error occurred"):
        return title, None, None
    
    transcript = get_video_transcript_with_timestamps(video_id)
    if isinstance(transcript, str):  # Error occurred
        return transcript, None, None
    
    safe_title = sanitize_filename(title)
    output_file_docx = f"{safe_title}.docx"
    output_file_pdf = f"{safe_title}.pdf"
    
    # Create Word document
    doc = Document()
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Add a blank line after the title
    
    # Create PDF
    pdf_buffer = io.BytesIO()
    pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
    styles.add(ParagraphStyle(name='Left', alignment=TA_LEFT))
    pdf_content = [Paragraph(title, styles['Center']), Spacer(1, 12)]
    
    for entry in transcript:
        start_time = int(entry['start'])
        minutes, seconds = divmod(start_time, 60)
        timestamp = f"{minutes}:{seconds:02d}"
        
        # Word document
        p = doc.add_paragraph()
        run = p.add_run(f"{timestamp}")
        run.bold = True
        p.add_run(f"\n{entry['text']}\n")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # PDF content
        pdf_content.append(Paragraph(f"<b>{timestamp}</b>", styles['Left']))
        pdf_content.append(Paragraph(entry['text'], styles['Left']))
        pdf_content.append(Spacer(1, 6))
    
    doc.save(output_file_docx)
    pdf.build(pdf_content)
    
    return "Transcript saved", output_file_docx, pdf_buffer

# Streamlit app
st.markdown('<div class="title-container"><h1>YouTube Transcript</h1></div>', unsafe_allow_html=True)

video_id = st.text_input("", key="video_id_input", placeholder="Enter YouTube Video ID")

if video_id:
    title, thumbnail_url = get_video_info(video_id)
    if thumbnail_url:
        st.markdown(f"""
        <div class="video-info">
            <img src="{thumbnail_url}" alt="Video Thumbnail">
            <span class="video-title">{title}</span>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.error(title)  # Display error message if video info couldn't be fetched

col1, col2, col3 = st.columns([1,2,1])
with col2:
    if st.button("Generate Transcript", key="generate_button"):
        if video_id:
            with st.spinner("Generating transcript..."):
                result, output_file_docx, pdf_buffer = save_transcript_to_word_and_pdf(video_id)
            
            if output_file_docx and pdf_buffer:
                st.success("Transcript generated successfully!")
                download_col1, download_col2 = st.columns(2)
                
                with download_col1:
                    with open(output_file_docx, "rb") as file:
                        docx_bytes = file.read()
                        st.download_button(
                            label="Download as Word",
                            data=docx_bytes,
                            file_name=os.path.basename(output_file_docx),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_word"
                        )
                
                with download_col2:
                    pdf_bytes = pdf_buffer.getvalue()
                    st.download_button(
                        label="Download as PDF",
                        data=pdf_bytes,
                        file_name=os.path.basename(output_file_docx).replace('.docx', '.pdf'),
                        mime="application/pdf",
                        key="download_pdf"
                    )
            else:
                st.error(result)
                st.error("Unable to generate transcript. Please check if the video has captions available.")
                st.markdown('<p class="medium-font">Possible reasons for this error:</p>', unsafe_allow_html=True)
                st.markdown("1. The video doesn't have any captions or transcripts.")
                st.markdown("2. The captions are disabled for this video.")
                st.markdown("3. The video ID is incorrect or the video doesn't exist.")
                st.markdown("4. There might be temporary issues with YouTube's transcript API.")
                st.markdown("Please try another video or check if the video ID is correct.")
        else:
            st.error("Please enter a YouTube Video ID.")

st.markdown("---")
st.markdown('<p class="big-font">Instructions:</p>', unsafe_allow_html=True)
st.markdown('<p class="medium-font">1. Enter the YouTube Video ID (e.g., \'dQw4w9WgXcQ\' from \'https://www.youtube.com/watch?v=dQw4w9WgXcQ\')</p>', unsafe_allow_html=True)
st.markdown('<p class="medium-font">2. Click \'Generate Transcript\' to create the Word and PDF documents</p>', unsafe_allow_html=True)
st.markdown('<p class="medium-font">3. Once generated, use the \'Download as Word\' or \'Download as PDF\' buttons to save the files</p>', unsafe_allow_html=True)
st.markdown('<p class="medium-font">Note: Not all videos have available transcripts. If you encounter an error, try another video.</p>', unsafe_allow_html=True)
